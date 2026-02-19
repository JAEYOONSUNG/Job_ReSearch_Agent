"""Generate a cover letter .docx from template by replacing placeholders."""

import difflib
import json
import os
import re
import sys
import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

RED = RGBColor(0xEE, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x70, 0xC0)
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(11)

TEMPLATE_PATH = Path(
    os.getenv(
        "CL_TEMPLATE_PATH",
        str(Path.home() / "Documents" / "CoverLetter_Template.docx"),
    )
)
OUTPUT_DIR = Path(
    os.getenv(
        "CL_OUTPUT_DIR",
        str(Path.home() / "Documents" / "CoverLetters"),
    )
)


def _set_run_format(run, color=RED):
    """Apply standard formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.color.rgb = color
    # Ensure East Asian font is also set for compatibility
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def _clear_paragraph(paragraph):
    """Remove all runs from a paragraph."""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def _replace_paragraph_text(paragraph, new_text, color=RED):
    """Replace all text in a paragraph with new_text in the given color."""
    _clear_paragraph(paragraph)
    run = paragraph.add_run(new_text)
    _set_run_format(run, color)


def _replace_in_runs(paragraph, old_text, new_text, color=RED):
    """Replace a placeholder string within paragraph runs, coloring only the replacement."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Collect run boundaries
    runs = paragraph.runs
    run_texts = [r.text for r in runs]
    combined = "".join(run_texts)
    idx = combined.find(old_text)
    if idx == -1:
        return False

    end_idx = idx + len(old_text)

    # Find which runs contain the placeholder
    char_pos = 0
    new_runs_data = []  # list of (text, color_override_or_None, source_run)

    for run in runs:
        rtext = run.text
        rstart = char_pos
        rend = char_pos + len(rtext)

        if rend <= idx or rstart >= end_idx:
            # Entirely outside placeholder
            new_runs_data.append((rtext, None, run))
        elif rstart >= idx and rend <= end_idx:
            # Entirely inside placeholder — will be replaced
            if rstart == idx:
                new_runs_data.append((new_text, color, run))
        else:
            # Partially overlapping
            if rstart < idx:
                # Part before placeholder
                new_runs_data.append((rtext[: idx - rstart], None, run))
                if rend <= end_idx:
                    new_runs_data.append((new_text, color, run))
                else:
                    new_runs_data.append((new_text, color, run))
                    new_runs_data.append((rtext[end_idx - rstart :], None, run))
            else:
                # rstart >= idx, rend > end_idx
                new_runs_data.append((rtext[end_idx - rstart :], None, run))

        char_pos = rend

    # Rebuild paragraph runs
    _clear_paragraph(paragraph)
    for text, clr, source_run in new_runs_data:
        if not text:
            continue
        new_run = paragraph.add_run(text)
        # Copy formatting from source run
        new_run.font.name = source_run.font.name or FONT_NAME
        new_run.font.size = source_run.font.size or FONT_SIZE
        new_run.bold = source_run.bold
        new_run.italic = source_run.italic
        if clr:
            new_run.font.color.rgb = clr
        elif source_run.font.color and source_run.font.color.rgb:
            new_run.font.color.rgb = source_run.font.color.rgb

    return True


def _insert_paragraph_after(paragraph, text, color=RED):
    """Insert a new paragraph after the given paragraph with styled text."""
    new_p = copy.deepcopy(paragraph._element)
    # Clear the copied paragraph's runs
    for child in list(new_p):
        if child.tag == qn("w:r"):
            new_p.remove(child)
    paragraph._element.addnext(new_p)

    # Get the new paragraph as a Paragraph object
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._element.getparent())

    run = new_para.add_run(text)
    _set_run_format(run, color)
    return new_para


def _insert_address_lines(doc, base_para_idx, address_lines):
    """Handle multi-line address by filling P08 and inserting additional paragraphs."""
    p08 = doc.paragraphs[base_para_idx]
    _replace_paragraph_text(p08, address_lines[0], RED)

    prev = p08
    for line in address_lines[1:]:
        prev = _insert_paragraph_after(prev, line, RED)


def generate_cover_letter(data: dict, output_path: str | None = None) -> str:
    """Generate a cover letter .docx from template and data dict.

    Args:
        data: Dict with keys: pi_name, institute, address, career_objective,
              topic_primary, topics_secondary, custom_paragraphs
        output_path: Optional output file path. Auto-generated if None.

    Returns:
        Path to the generated .docx file.
    """
    doc = Document(str(TEMPLATE_PATH))

    pi_name = data["pi_name"]
    institute = data["institute"]
    address = data["address"]
    career_objective = data["career_objective"]
    topic_primary = data["topic_primary"]
    topics_secondary = data["topics_secondary"]
    custom_paragraphs = data.get("custom_paragraphs", [])

    # --- P06: [name], Ph.D. ---
    _replace_paragraph_text(doc.paragraphs[6], f"{pi_name}, Ph.D.", RED)

    # --- P07: [institute] ---
    _replace_paragraph_text(doc.paragraphs[7], institute, RED)

    # --- P08: [address] (may be multi-line) ---
    address_lines = [line.strip() for line in address.split("\n") if line.strip()]
    if len(address_lines) <= 1:
        _replace_paragraph_text(doc.paragraphs[8], address_lines[0] if address_lines else address, RED)
    else:
        _insert_address_lines(doc, 8, address_lines)

    # After address insertion, paragraph indices may have shifted.
    # Re-locate key paragraphs by content matching.
    def _find_para(prefix):
        for i, p in enumerate(doc.paragraphs):
            if p.text.startswith(prefix):
                return i
        return None

    # --- P10: Dear Prof. [name], ---
    dear_idx = _find_para("Dear Prof.")
    if dear_idx is not None:
        _replace_paragraph_text(doc.paragraphs[dear_idx], f"Dear Prof. {pi_name},\n", RED)

    # --- P11: career objective ---
    # Rebuild the intro paragraph: keep static prefix in black, insert career objective in red
    intro_idx = _find_para("I am writing to apply")
    if intro_idx is not None:
        p = doc.paragraphs[intro_idx]
        _clear_paragraph(p)
        prefix = "I am writing to apply for the postdoctoral position in your laboratory. My career objective is "
        r1 = p.add_run(prefix)
        r1.font.name = FONT_NAME
        r1.font.size = FONT_SIZE
        r2 = p.add_run(career_objective)
        _set_run_format(r2, RED)
        r3 = p.add_run(".")
        r3.font.name = FONT_NAME
        r3.font.size = FONT_SIZE

    # --- P16: topics ---
    topic_idx = _find_para("I have closely followed")
    if topic_idx is not None:
        _replace_in_runs(doc.paragraphs[topic_idx], "[topic 1]", topic_primary, RED)
        _replace_in_runs(doc.paragraphs[topic_idx], "[topic2, 3, 4]", topics_secondary, RED)

        # Insert custom paragraphs after the topic paragraph
        prev = doc.paragraphs[topic_idx]
        for cp_text in custom_paragraphs:
            prev = _insert_paragraph_after(prev, cp_text, RED)

    # --- Determine output path ---
    if output_path is None:
        label = data.get("label", pi_name)
        filename = f"Cover Letter for postdoctoral application_{label}.docx"
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        output_path = str(OUTPUT_DIR / filename)

    doc.save(output_path)
    return output_path


def _normalize_quotes(text: str) -> str:
    """Normalize smart/curly quotes to straight quotes for comparison."""
    return text.replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')


def _tokenize(text: str) -> list[str]:
    """Split text into tokens preserving whitespace as separate tokens."""
    return re.findall(r'\S+|\s+', text)


def _build_char_color_map(paragraph) -> list:
    """Build per-character color map from a paragraph's runs."""
    colors = []
    for run in paragraph.runs:
        c = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        colors.extend([c] * len(run.text))
    return colors


def _diff_replace_paragraph(paragraph, new_full_text, highlight_color=BLUE):
    """Replace paragraph via word-level diff; only changed tokens get highlight_color."""
    old_text = paragraph.text

    if old_text == new_full_text:
        return  # Nothing changed

    # Build per-character color map from current paragraph
    char_colors = _build_char_color_map(paragraph)

    old_tokens = _tokenize(old_text)
    new_tokens = _tokenize(new_full_text)

    # Normalize quotes for comparison only; output uses original new_tokens.
    old_norm = [_normalize_quotes(t) for t in old_tokens]
    new_norm = [_normalize_quotes(t) for t in new_tokens]

    sm = difflib.SequenceMatcher(None, old_norm, new_norm, autojunk=False)

    def _split_equal_at_color_boundaries(chunk, old_char_pos):
        """Split an equal chunk into sub-segments at color boundaries."""
        result = []
        current_color = char_colors[old_char_pos] if old_char_pos < len(char_colors) else None
        seg_start = 0
        for ci in range(len(chunk)):
            c_pos = old_char_pos + ci
            c_color = char_colors[c_pos] if c_pos < len(char_colors) else None
            if c_color != current_color:
                if ci > seg_start:
                    result.append((chunk[seg_start:ci], current_color))
                current_color = c_color
                seg_start = ci
        if seg_start < len(chunk):
            result.append((chunk[seg_start:], current_color))
        return result

    _DEL = object()  # Sentinel for deletion boundaries

    def _process_opcodes(matcher, old_toks, new_toks, old_toks_orig, new_toks_orig, char_base, depth=0):
        """Process opcodes and return segments with _DEL markers for deletions."""
        result = []
        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                chunk = ''.join(new_toks_orig[j1:j2])
                ocp = char_base + len(''.join(old_toks_orig[:i1]))
                result.extend(_split_equal_at_color_boundaries(chunk, ocp))
            elif op == 'replace':
                if depth < 1:
                    old_sub = old_toks[i1:i2]
                    new_sub = new_toks[j1:j2]
                    old_sub_orig = old_toks_orig[i1:i2]
                    new_sub_orig = new_toks_orig[j1:j2]
                    sub_sm = difflib.SequenceMatcher(None, old_sub, new_sub, autojunk=False)
                    sub_base = char_base + len(''.join(old_toks_orig[:i1]))
                    result.extend(_process_opcodes(
                        sub_sm, old_sub, new_sub, old_sub_orig, new_sub_orig, sub_base, depth + 1))
                else:
                    chunk = ''.join(new_toks_orig[j1:j2])
                    result.append((chunk, highlight_color))
            elif op == 'insert':
                chunk = ''.join(new_toks_orig[j1:j2])
                result.append((chunk, highlight_color))
            elif op == 'delete':
                result.append((_DEL, None))
        return result

    raw_segments = _process_opcodes(sm, old_norm, new_norm, old_tokens, new_tokens, 0)

    # Post-process: at each _DEL marker, find the nearest space and color it
    # as highlight to signal deletion boundaries.
    segments = []
    for seg in raw_segments:
        if seg[0] is _DEL:
            # Look back: if previous segment ends with space, recolor it
            if segments and segments[-1][0] and segments[-1][0][-1] == ' ':
                prev_text, prev_color = segments[-1]
                if prev_color != highlight_color:
                    segments[-1] = (prev_text[:-1], prev_color)
                    segments.append((' ', highlight_color))
            # Otherwise the next segment's leading space will be handled below
            else:
                segments.append((_DEL, None))  # keep marker for forward scan
        else:
            # Check if previous item is a leftover _DEL marker
            if segments and segments[-1][0] is _DEL:
                segments.pop()  # remove marker
                text, color = seg
                if text and text[0] == ' ' and color != highlight_color:
                    segments.append((' ', highlight_color))
                    seg = (text[1:], color)
                    if not seg[0]:
                        continue
            segments.append(seg)
    # Remove any remaining _DEL markers
    segments = [(t, c) for t, c in segments if t is not _DEL]

    # Absorb whitespace-only segments into surrounding colored segments.
    for i in range(len(segments)):
        text, color = segments[i]
        if text.isspace() and color != highlight_color:
            prev_color = segments[i - 1][1] if i > 0 else None
            next_color = segments[i + 1][1] if i + 1 < len(segments) else None
            if prev_color == next_color and prev_color is not None:
                segments[i] = (text, prev_color)
            elif next_color is not None:
                segments[i] = (text, next_color)
            elif prev_color is not None:
                segments[i] = (text, prev_color)

    # Merge adjacent segments with the same color to reduce run count
    merged = []
    for text, color in segments:
        if merged and merged[-1][1] == color:
            merged[-1] = (merged[-1][0] + text, color)
        else:
            merged.append((text, color))

    # Rebuild paragraph
    _clear_paragraph(paragraph)
    for text, color in merged:
        if not text:
            continue
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        if color:
            run.font.color.rgb = color


def refine_cover_letter(data: dict) -> str:
    """Refine an existing draft cover letter, marking only changed words in blue.

    Args:
        data: Dict with keys:
            - input_path: path to the draft .docx to refine
            - replacements: list of {paragraph_index, new_text, full_replace?}
            - output_path: (optional) output file path

    Returns:
        Path to the refined .docx file.
    """
    input_path = data["input_path"]
    replacements = data["replacements"]

    doc = Document(input_path)
    orig_doc = Document(input_path)

    for repl in replacements:
        idx = repl["paragraph_index"]
        new_text = repl["new_text"]
        full_replace = repl.get("full_replace", False)

        if idx < 0 or idx >= len(doc.paragraphs):
            continue

        para = doc.paragraphs[idx]
        orig_para = orig_doc.paragraphs[idx]

        if full_replace:
            # new_text is the complete paragraph — diff against old full text
            _diff_replace_paragraph(para, new_text, BLUE)
        else:
            # new_text replaces only the colored (red/blue) portion.
            # Reconstruct the full new text: black_prefix + new_text + trailing
            black_prefix_parts = []
            for r in orig_para.runs:
                is_colored = r.font.color and r.font.color.rgb in (RED, BLUE)
                if not is_colored:
                    black_prefix_parts.append(r.text)
                else:
                    break

            trailing_parts = []
            found_colored = False
            for r in orig_para.runs:
                is_colored = r.font.color and r.font.color.rgb in (RED, BLUE)
                if is_colored:
                    found_colored = True
                elif found_colored:
                    trailing_parts.append(r.text)

            full_new = ''.join(black_prefix_parts) + new_text + ''.join(trailing_parts)
            _diff_replace_paragraph(para, full_new, BLUE)

    # Determine output path
    output_path = data.get("output_path")
    if output_path is None:
        p = Path(input_path)
        output_path = str(p.parent / f"{p.stem}_refined{p.suffix}")

    doc.save(output_path)
    return output_path


def main():
    """CLI entry point: reads JSON from stdin or file argument.

    Usage:
        python -m src.coverletter.docx_writer <json_file>           # generate
        python -m src.coverletter.docx_writer --refine <json_file>  # refine
    """
    if len(sys.argv) > 1 and sys.argv[1] == "--refine":
        json_path = sys.argv[2] if len(sys.argv) > 2 else None
        if json_path:
            with open(json_path) as f:
                data = json.load(f)
        else:
            data = json.load(sys.stdin)
        result_path = refine_cover_letter(data)
        print(result_path)
    else:
        if len(sys.argv) > 1:
            json_path = sys.argv[1]
            with open(json_path) as f:
                data = json.load(f)
        else:
            data = json.load(sys.stdin)
        output = data.get("output_path")
        result_path = generate_cover_letter(data, output)
        print(result_path)


if __name__ == "__main__":
    main()
